import io
import os
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image

from app.services.docx_service import export_product_docx, import_docx, parse_markdown_blocks, _set_table_widths
from app.services.skill_service import build_skill_prompt, discover_skills, safe_extract_zip
from app.routers.workspace import _user_storage_key
from app.services.memorybear import build_memory_context
from app.services.generation_task_service import build_generation_meta
from app.models import GenerationTask, InstalledSkill
from app.services.generation_task_service import _exception_detail


def make_docx_with_table_and_image():
    image_stream = io.BytesIO()
    Image.new("RGB", (80, 50), "#2b8a68").save(image_stream, "PNG")
    image_stream.seek(0)
    document = Document()
    document.add_heading("图文笔记", level=1)
    document.add_paragraph("正文第一段")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "分块生成"
    table.cell(1, 1).text = "提升长文质量"
    document.add_picture(image_stream)
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_docx_import_preserves_heading_table_and_image(tmp_path):
    saved = []
    def save(name, media_type, blob, anchor):
        target = tmp_path / name
        target.write_bytes(blob)
        saved.append((target, anchor))
        return f"/api/assets/{len(saved)}"
    result = import_docx(make_docx_with_table_and_image(), save)
    assert "<h1>图文笔记</h1>" in result["html"]
    assert "<table>" in result["html"]
    assert "分块生成" in result["plain_text"]
    assert "[插图 1" in result["plain_text"]
    assert len(saved) == 1 and saved[0][0].is_file()


def test_word_export_embeds_source_image(tmp_path):
    image = tmp_path / "source.png"
    Image.new("RGB", (60, 40), "#137c5a").save(image)
    data = export_product_docx(
        {"title": "产品", "product_type": "sop", "content": "# 方案\n\n[插图 1: source.png]\n\n## 验收", "price_suggestion": 99},
        None,
        [{"filename": "source.png", "storage_path": str(image)}],
    )
    parsed = Document(io.BytesIO(data))
    assert parsed.inline_shapes
    assert "产品" in "\n".join(p.text for p in parsed.paragraphs)


def test_word_table_widths_never_go_negative_for_wide_tables():
    document = Document()
    table = document.add_table(rows=1, cols=12)
    _set_table_widths(table, [f"列{i}" for i in range(12)], [["值"] * 12])
    widths = [int(col.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for col in table._tbl.tblGrid.gridCol_lst]
    assert sum(widths) == int(6.3 * 914400)
    assert min(widths) > 0


def test_skill_zip_discovery_and_prompt(tmp_path):
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("course-maker/SKILL.md", "---\nname: course-maker\ndescription: 做课质量流程\n---\n# 工作流\n先规划再验收")
    safe_extract_zip(stream.getvalue(), tmp_path)
    skills = discover_skills(tmp_path)
    assert skills[0]["name"] == "course-maker"
    assert "先规划再验收" in build_skill_prompt(skills)


def test_skill_product_labels_use_stable_zero_to_thirteen_ids():
    from app.services.skill_service import product_type_ids_for_skill
    assert 1 in product_type_ids_for_skill("guizang-ppt-skill")
    assert all(0 <= value <= 13 for value in product_type_ids_for_skill("wechat-article-pro"))
    assert 12 in product_type_ids_for_skill("promo-storyboard", "video shot planning")
    assert product_type_ids_for_skill("unclassified-helper") == [13]


def test_v3_python_memorybear_source_is_vendored():
    from app.services.memorybear_python_adapter import VENDOR_ROOT
    assert (VENDOR_ROOT / "memory_bear" / "engine.py").is_file()
    assert (VENDOR_ROOT / "pyproject.toml").is_file()


def test_skill_zip_rejects_path_traversal(tmp_path):
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr("../outside/SKILL.md", "unsafe")
    try:
        safe_extract_zip(stream.getvalue(), tmp_path)
        assert False, "path traversal should fail"
    except ValueError as exc:
        assert "不安全路径" in str(exc)


def test_skill_listing_never_installs_or_scans_archives(monkeypatch):
    from app.routers import workspace

    class Skills:
        def list(self):
            return [{"id": 1, "name": "fast-list", "description": "只读列表", "category": "测试"}]

    monkeypatch.setattr(workspace, "table", lambda name, user: Skills())
    monkeypatch.setattr(workspace, "seed_bundled_skills", lambda user: (_ for _ in ()).throw(AssertionError("GET 不得安装 Skill")))
    rows = workspace.list_skills(user={"id": "local:test"}, q=None, category=None, limit=30, offset=0)
    assert [item["name"] for item in rows] == ["fast-list"]


def test_task_strategy_endpoint_exposes_defaults(client):
    response = client.get("/api/tasks/strategies")
    assert response.status_code == 200
    data = response.json()
    assert any(item["id"] == "chunked_generation" for item in data["algorithms"])
    assert "article" in data["defaults"]


def test_background_task_creation_persists_strategy(client, sample_note, monkeypatch):
    monkeypatch.setattr("app.routers.tasks.enqueue_generation", lambda task_id, user: None)
    response = client.post("/api/tasks", json={
        "note_id": sample_note.id,
        "product_types": ["sop"],
        "skill_ids": [],
        "algorithms": ["chunked_generation"],
        "techniques": ["source_grounding", "hallucination_check"],
    })
    assert response.status_code == 200
    task = response.json()
    assert task["status"] == "queued"
    assert task["algorithms"] == ["chunked_generation"]
    assert client.get(f"/api/tasks/{task['id']}").json()["product_types"] == ["sop"]


def test_task_response_contains_subject_and_retry_copies_configuration(client, sample_note, db_session, monkeypatch):
    queued = []
    monkeypatch.setattr("app.routers.tasks.enqueue_generation", lambda task_id, user: queued.append(task_id))
    original = client.post("/api/tasks", json={
        "note_id": sample_note.id,
        "product_types": ["ppt"],
        "skill_ids": [11],
        "algorithms": ["iterative_refinement"],
        "techniques": ["quality_scoring"],
        "product_strategies": {"ppt": {"skill_ids": [12]}},
    }).json()
    assert original["subject_id"] == sample_note.subject_id
    assert original["subject_name"]

    stored = db_session.get(GenerationTask, original["id"])
    stored.status = "failed"
    db_session.commit()
    retried = client.post(f"/api/tasks/{original['id']}/retry")
    assert retried.status_code == 200
    task = retried.json()
    assert task["id"] != original["id"]
    assert task["skill_ids"] == [11]
    assert task["algorithms"] == ["iterative_refinement"]
    assert task["product_strategies"] == {"ppt": {"skill_ids": [12]}}
    assert queued[-1] == task["id"]


def test_iterative_refinement_is_only_an_algorithm(client):
    strategies = client.get("/api/tasks/strategies").json()
    assert any(item["id"] == "iterative_refinement" for item in strategies["algorithms"])
    assert all(item["id"] != "iterative_refinement" for item in strategies["techniques"])


def test_task_with_per_product_strategy_persists(client, sample_note, monkeypatch):
    """2026-08 feat/29：创建任务时附带 product_strategies，按 product_type 独立指定 strategy。"""
    monkeypatch.setattr("app.routers.tasks.enqueue_generation", lambda task_id, user: None)
    response = client.post("/api/tasks", json={
        "note_id": sample_note.id,
        "product_types": ["article", "ppt"],
        "skill_ids": [1, 2],
        "algorithms": ["hierarchical_planning"],
        "techniques": ["source_grounding", "quality_scoring"],
        "product_strategies": {
            "article": {"skill_ids": [99], "algorithms": ["chunked_generation"], "techniques": ["source_grounding"]},
            "ppt": {"algorithms": ["iterative_refinement"]},
        },
    })
    assert response.status_code == 200
    task = response.json()
    assert task["product_strategies"]["article"]["skill_ids"] == [99]
    assert task["product_strategies"]["ppt"]["algorithms"] == ["iterative_refinement"]
    # 详情接口也能取到
    detail = client.get(f"/api/tasks/{task['id']}").json()
    assert detail["product_strategies"]["article"]["techniques"] == ["source_grounding"]


def test_task_rejects_product_strategies_with_unknown_product_type(client, sample_note, monkeypatch):
    """product_strategies 中的 key 必须在 product_types 内，否则 400。"""
    monkeypatch.setattr("app.routers.tasks.enqueue_generation", lambda task_id, user: None)
    response = client.post("/api/tasks", json={
        "note_id": sample_note.id,
        "product_types": ["article"],
        "product_strategies": {
            "article": {"skill_ids": []},
            "sop": {"skill_ids": []},  # 不在 product_types 中
        },
    })
    assert response.status_code == 400
    assert "sop" in response.json()["detail"]


def test_terminal_generation_task_can_be_deleted_but_running_task_is_protected(client, sample_note, db_session, monkeypatch):
    monkeypatch.setattr("app.routers.tasks.enqueue_generation", lambda task_id, user: None)
    task = client.post("/api/tasks", json={
        "note_id": sample_note.id,
        "product_types": ["article"],
    }).json()

    running = client.delete(f"/api/tasks/{task['id']}")
    assert running.status_code == 409

    stored = db_session.get(GenerationTask, task["id"])
    stored.status = "failed"
    db_session.commit()
    deleted = client.delete(f"/api/tasks/{task['id']}")
    assert deleted.status_code == 200
    assert deleted.json() == {"ok": True}
    assert client.get(f"/api/tasks/{task['id']}").status_code == 404


def test_generation_error_detail_keeps_exception_type_when_message_is_empty():
    assert _exception_detail(TimeoutError()) == "TimeoutError: 无错误详情"


def test_user_storage_key_is_windows_safe_and_stable():
    user = {"id": "local:editor@example.com"}
    key = _user_storage_key(user)
    assert key == _user_storage_key(user)
    assert len(key) == 24
    assert key.isalnum()


def test_memorybear_context_has_layers_and_prunes_irrelevant_history():
    context, meta = build_memory_context(
        {"id": 1, "title": "RAG 课程", "raw_content": "向量检索和召回评估"},
        {"name": "人工智能"},
        [{"id": 2, "title": "RAG 实践", "raw_content": "召回率和向量数据库", "created_at": "2026-08-04T00:00:00+00:00"},
         {"id": 3, "title": "无关记录", "raw_content": "烘焙配方", "created_at": "2026-08-04T00:00:00+00:00"}],
        [{"title": "RAG 产品", "content": "向量检索课程产品", "created_at": "2026-08-04T00:00:00+00:00"}],
    )
    # 五层结构 + 相关历史保留
    assert "工作记忆" in context and "RAG 实践" in context
    assert meta["layers"]["episodic"] >= 1
    # 无关记录（烘焙配方）因零相关性被激活度阈值过滤，不进入上下文
    assert "烘焙配方" not in context
    # 场景路由：历史充足 → MemoryBear 主导
    assert meta["scene_router"]["memorybear_weight"] > meta["scene_router"]["rag_weight"]


def test_generation_meta_links_all_products_to_source_task():
    # 🔍 [作用] 模块级回归：产品溯源必须包含任务编号与策略信息。
    meta = build_generation_meta(
        42,
        {"skill_ids": [7], "algorithms": ["hierarchical_planning"]},
        [{"id": 7, "name": "interview-bank"}],
        ["memorybear", "rag_grounding"],
        {"layers": {"working": 1}},
        ["示例警告"],
    )
    assert meta["task_id"] == 42
    assert meta["skill_names"] == ["interview-bank"]
    assert meta["techniques"] == ["memorybear", "rag_grounding"]


# =============================================================================
# 2026-08 feat/28：Skills 批量上传
# =============================================================================
def _make_skill_zip(skill_name: str, description: str = "示例 Skill") -> bytes:
    """构造一个含 SKILL.md 的最小 zip 字节流。"""
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        archive.writestr(
            f"{skill_name}/SKILL.md",
            f"---\nname: {skill_name}\ndescription: {description}\n---\n# 工作流\n执行 {skill_name} 任务",
        )
    return stream.getvalue()


def test_batch_upload_skills_installs_each_zip(client):
    """一次性上传多个 zip，每个 zip 内的 SKILL.md 各自入库。"""
    zip_a = _make_skill_zip("course-maker")
    zip_b = _make_skill_zip("ppt-designer", "PPT 美化指南")
    response = client.post(
        "/api/skills/batch-upload",
        files=[
            ("files", ("a.zip", zip_a, "application/zip")),
            ("files", ("b.zip", zip_b, "application/zip")),
        ],
        data={"category": "知识产品"},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["received"] == 2
    assert data["installed"] == 2
    assert data["success"] is True
    assert data["failures"] == []
    assert {s["name"] for s in data["skills"]} == {"course-maker", "ppt-designer"}


def test_batch_upload_skips_duplicate_skill_and_reports_its_name(client):
    archive = _make_skill_zip("duplicate-skill")
    first = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("first.zip", archive, "application/zip"))],
    )
    second = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("second.zip", archive, "application/zip"))],
    )

    assert first.status_code == second.status_code == 200
    assert first.json()["installed"] == 1
    assert second.json()["installed"] == 0
    assert second.json()["duplicates"] == ["duplicate-skill"]
    assert second.json()["per_archive"][0]["duplicates"] == ["duplicate-skill"]
    listed = client.get("/api/skills").json()
    assert [item["name"] for item in listed].count("duplicate-skill") == 1


def test_skill_list_hides_historical_duplicates(client, db_session):
    first = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("case.zip", _make_skill_zip("Case Skill"), "application/zip"))],
    ).json()["skills"][0]
    second = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("other.zip", _make_skill_zip("Other Skill"), "application/zip"))],
    ).json()["skills"][0]
    stored_second = db_session.get(InstalledSkill, second["id"])
    assert stored_second is not None
    stored_second.name = "  CASE   SKILL  "
    db_session.commit()

    listed = client.get("/api/skills").json()
    matching = [item for item in listed if " ".join(item["name"].split()).casefold() == "case skill"]
    assert len(matching) == 1
    assert matching[0]["id"] in {first["id"], second["id"]}


def test_skill_list_supports_bounded_lazy_pagination(client):
    for name in ("lazy-page-one", "lazy-page-two", "lazy-page-three"):
        response = client.post(
            "/api/skills/batch-upload",
            files=[("files", (f"{name}.zip", _make_skill_zip(name), "application/zip"))],
        )
        assert response.status_code == 200

    first = client.get("/api/skills", params={"q": "lazy-page", "limit": 2, "offset": 0})
    second = client.get("/api/skills", params={"q": "lazy-page", "limit": 2, "offset": 2})

    assert first.status_code == second.status_code == 200
    assert len(first.json()) == 2
    assert len(second.json()) == 1
    assert {item["id"] for item in first.json()}.isdisjoint({item["id"] for item in second.json()})
    assert all(item.get("product_type_ids") for item in first.json() + second.json())


def test_batch_upload_installs_multiple_skills_from_one_bundle(client):
    """A single bundle ZIP may contain several independent Skill folders."""
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w") as archive:
        for name in ("bundle-course", "bundle-ppt", "bundle-checklist"):
            archive.writestr(
                f"skills/{name}/SKILL.md",
                f"---\nname: {name}\ndescription: bundled test\n---\n# {name}\n",
            )

    response = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("all-skills.zip", stream.getvalue(), "application/zip"))],
    )

    assert response.status_code == 200
    data = response.json()
    assert data["received"] == 1
    assert data["installed"] == 3
    assert {skill["name"] for skill in data["skills"]} == {
        "bundle-course", "bundle-ppt", "bundle-checklist",
    }


def test_failed_skill_upload_never_changes_subjects_or_notes(client, sample_subject, sample_note):
    before_subjects = client.get("/api/subjects").json()
    before_notes = client.get("/api/notes").json()

    response = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("broken.zip", b"not-a-zip", "application/zip"))],
    )

    assert response.status_code == 200
    assert response.json()["success"] is False
    assert client.get("/api/subjects").json() == before_subjects
    assert client.get("/api/notes").json() == before_notes
    subject = next(item for item in before_subjects if item["id"] == sample_subject.id)
    assert subject["note_count"] == 1


def test_creating_duplicate_subject_uses_stable_numeric_suffix(client):
    first = client.post("/api/subjects", json={"name": "重复科目"})
    second = client.post("/api/subjects", json={"name": "重复科目"})
    third = client.post("/api/subjects", json={"name": "重复科目"})

    assert first.status_code == second.status_code == third.status_code == 200
    assert [first.json()["name"], second.json()["name"], third.json()["name"]] == [
        "重复科目", "重复科目-1", "重复科目-2",
    ]


def test_subject_suffixes_skip_existing_names_and_apply_to_rename(client):
    base = client.post("/api/subjects", json={"name": "名称"}).json()
    reserved = client.post("/api/subjects", json={"name": "名称-1"}).json()
    duplicate = client.post("/api/subjects", json={"name": "名称"}).json()
    other = client.post("/api/subjects", json={"name": "其他"}).json()

    renamed = client.put(f"/api/subjects/{other['id']}", json={"name": "名称"})

    assert reserved["name"] == "名称-1"
    assert duplicate["name"] == "名称-2"
    assert renamed.status_code == 200
    assert renamed.json()["name"] == "名称-3"
    assert client.get(f"/api/subjects/{base['id']}").json()["name"] == "名称"


def test_batch_upload_accepts_skill_archive_larger_than_20_mb(client):
    stream = io.BytesIO()
    with zipfile.ZipFile(stream, "w", compression=zipfile.ZIP_STORED) as archive:
        archive.writestr(
            "large-skill/SKILL.md",
            "---\nname: large-skill\ndescription: 大包回归测试\n---\n# 工作流\n验证大包安装",
        )
        archive.writestr("large-skill/assets/payload.bin", os.urandom(20 * 1024 * 1024 + 1))
    assert len(stream.getvalue()) > 20 * 1024 * 1024

    response = client.post(
        "/api/skills/batch-upload",
        files=[("files", ("large-skill.zip", stream.getvalue(), "application/zip"))],
    )

    assert response.status_code == 200
    assert response.json()["installed"] == 1


def test_batch_upload_skills_rejects_non_zip_files(client):
    """所有文件都不是 zip → 400 错误，明确列出非 zip 文件名。"""
    response = client.post(
        "/api/skills/batch-upload",
        files=[
            ("files", ("a.txt", b"hello", "text/plain")),
            ("files", ("b.png", b"\x89PNG", "image/png")),
        ],
    )
    assert response.status_code == 400
    assert "a.txt" in response.json()["detail"]


def test_batch_upload_skills_partial_failure_is_reported(client):
    """部分文件无效时，成功的仍正常安装，失败的在 failures 里报告。"""
    zip_a = _make_skill_zip("ok-skill")
    bad_zip = b"not a real zip"
    response = client.post(
        "/api/skills/batch-upload",
        files=[
            ("files", ("ok.zip", zip_a, "application/zip")),
            ("files", ("bad.zip", bad_zip, "application/zip")),
        ],
    )
    assert response.status_code == 200
    data = response.json()
    assert data["received"] == 2
    assert data["installed"] == 1
    assert data["success"] is False
    assert len(data["failures"]) == 1
    assert data["failures"][0]["archive"] == "bad.zip"
    assert data["skills"][0]["name"] == "ok-skill"


def test_batch_upload_skills_empty_files_returns_400(client):
    """未传 files → 400（业务校验）或 422（FastAPI 字段必填校验）；任一均可接受。"""
    response = client.post("/api/skills/batch-upload", files=[])
    assert response.status_code in (400, 422)
