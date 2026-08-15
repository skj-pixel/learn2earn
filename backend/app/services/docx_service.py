"""DOCX import/export with ordered text, tables and embedded images."""
from __future__ import annotations

import html
import io
import re
from pathlib import Path
from typing import Callable

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

# 正文插图标记，如 "[插图 1: 文件名，位置 block-3]"
IMAGE_MARKER = re.compile(r"\[插图\s*(\d+)[^\]]*\]")


def iter_blocks(document: DocumentType):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def import_docx(data: bytes, save_image: Callable[[str, str, bytes, str], str]) -> dict:
    doc = Document(io.BytesIO(data))
    html_parts, plain_parts, assets = [], [], []
    image_index = 0
    for block_index, block in enumerate(iter_blocks(doc)):
        if isinstance(block, Table):
            rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            if not rows:
                continue
            html_parts.append('<table><tbody>')
            for row_index, row in enumerate(rows):
                tag = "th" if row_index == 0 else "td"
                html_parts.append("<tr>" + "".join(f"<{tag}>{html.escape(value)}</{tag}>" for value in row) + "</tr>")
            html_parts.append("</tbody></table>")
            plain_parts.extend(" | ".join(row) for row in rows)
            continue

        text = block.text.strip()
        style = (block.style.name or "") if block.style else ""
        if text:
            match = re.search(r"Heading\s*([1-6])", style, re.I)
            if match:
                level = match.group(1)
                html_parts.append(f"<h{level}>{html.escape(text)}</h{level}>")
            elif "List" in style:
                html_parts.append(f"<p>• {html.escape(text)}</p>")
            else:
                html_parts.append(f"<p>{html.escape(text)}</p>")
            plain_parts.append(text)

        for run in block.runs:
            for drawing in run._element.xpath('.//a:blip'):
                rel_id = drawing.get(qn('r:embed'))
                if not rel_id or rel_id not in doc.part.related_parts:
                    continue
                image_index += 1
                part = doc.part.related_parts[rel_id]
                suffix = Path(str(part.partname)).suffix or ".png"
                name = f"word-image-{image_index}{suffix}"
                anchor = f"block-{block_index}"
                url = save_image(name, part.content_type, part.blob, anchor)
                html_parts.append(f'<figure data-source-anchor="{anchor}"><img src="{html.escape(url)}" alt="Word 插图 {image_index}"><figcaption>插图 {image_index}</figcaption></figure>')
                assets.append({"filename": name, "url": url, "source_anchor": anchor})
                plain_parts.append(f"[插图 {image_index}: {name}，位置 {anchor}]")
    return {"html": "\n".join(html_parts), "plain_text": "\n\n".join(plain_parts), "assets": assets}


# ============================================================
# Markdown -> 结构化块（支持表格/代码/有序列表/行内格式）
# ============================================================

def _split_row(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [c.strip() for c in row.split("|")]


def _is_special_start(line: str) -> bool:
    return bool(
        line.startswith("#") or line.startswith(">") or line.startswith("```")
        or re.match(r"^[-*+]\s+", line)
        or re.match(r"^\d+[\.\)]\s+", line)
        or re.match(r"^(\*\*\*|---|___)\s*$", line)
    )


def _inline_segments(text: str) -> list[tuple[str, dict]]:
    """把行内 Markdown（**粗体** *斜体* `代码` ~~删除线~~）切成带样式标记的片段。"""
    pattern = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|~~.+?~~)")
    segments: list[tuple[str, dict]] = []
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], {}))
        tok = m.group(0)
        if tok.startswith("**"):
            segments.append((tok[2:-2], {"bold": True}))
        elif tok.startswith("`"):
            segments.append((tok[1:-1], {"code": True}))
        elif tok.startswith("~~"):
            segments.append((tok[2:-2], {"strike": True}))
        else:
            segments.append((tok[1:-1], {"italic": True}))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], {}))
    return [(t, f) for t, f in segments if t]


def _add_runs(paragraph, text: str):
    for seg, flags in _inline_segments(text):
        run = paragraph.add_run(seg)
        if flags.get("bold"):
            run.bold = True
        if flags.get("italic"):
            run.italic = True
        if flags.get("strike"):
            run.font.strike = True
        if flags.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def _strip_markers(text: str) -> str:
    return IMAGE_MARKER.sub("", text).strip()


def parse_markdown_blocks(content: str) -> list[tuple]:
    """把产品正文（Markdown）解析为结构化块，供 export_product_docx 使用。"""
    lines = (content or "").splitlines()
    blocks: list[tuple] = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # 代码块
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过收尾 ```
            blocks.append(("code", lang, "\n".join(code_lines)))
            continue
        # 分隔线
        if re.match(r"^(\*\*\*|---|___)\s*$", stripped):
            blocks.append(("hr",))
            i += 1
            continue
        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        # 表格：当前行含 | 且下一行是分隔行
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", lines[i + 1]) and "-" in lines[i + 1]:
            header = _split_row(stripped)
            i += 2  # 跳过表头 + 分隔行
            rows = []
            while i < n and lines[i].strip() and "|" in lines[i].strip():
                rows.append(_split_row(lines[i].strip()))
                i += 1
            blocks.append(("table", header, rows))
            continue
        # 引用
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append(("quote", " ".join(quote_lines)))
            continue
        # 有序列表
        if re.match(r"^\d+[\.\)]\s+", stripped):
            items = []
            while i < n and re.match(r"^\d+[\.\)]\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+[\.\)]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ordered", items))
            continue
        # 无序列表
        if re.match(r"^[-*+]\s+", stripped):
            items = []
            while i < n and re.match(r"^[-*+]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*+]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("bullet", items))
            continue
        # 段落：聚合连续普通行
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _is_special_start(lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(("p", " ".join(para_lines)))
    return blocks


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_table_widths(table: Table, header: list[str], rows: list[list[str]], total_width: Emu = Inches(6.3)):
    """给导出 Word 的表格设置固定列宽（按每列内容长度成比例分配），避免 Word 自动拉伸导致排版错乱。

    - 固定布局（w:tblLayout type=fixed）确保列宽被严格遵守
    - 同步设置 tblGrid 列宽与每个单元格 tcW，保证窄列不被压扁、宽列不被换行挤乱
    - 表头行跨页重复
    """
    ncols = max(1, len(header))
    col_lens = [len(str(header[c])) if c < len(header) else 0 for c in range(ncols)]
    for row in rows:
        for c in range(ncols):
            val = row[c] if c < len(row) else ""
            col_lens[c] = max(col_lens[c], len(str(val)))
    total = max(1, sum(col_lens))
    total_int = int(total_width)
    # 每列至少 600 twips，但先按总宽分配余量，避免多列时最后一列变成负宽度。
    minimum = 600
    if ncols * minimum > total_int:
        minimum = max(240, total_int // ncols)
    remaining = total_int - minimum * ncols
    weights = [max(1, length) for length in col_lens]
    weight_total = sum(weights)
    widths = [minimum + (remaining * weight // weight_total) for weight in weights]
    widths[-1] += total_int - sum(widths)

    # 固定布局：让列宽设置生效
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    # 网格列宽（覆盖默认自动列宽）
    tblGrid = table._tbl.tblGrid
    grid_cols = tblGrid.findall(qn("w:gridCol"))
    for c, w in enumerate(widths):
        if c < len(grid_cols):
            grid_cols[c].set(qn("w:w"), str(w))
        else:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(w))
            tblGrid.append(gc)

    def _set_cell_width(cell, w):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(w))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    for row in table.rows:
        for c, cell in enumerate(row.cells):
            _set_cell_width(cell, widths[min(c, ncols - 1)])

    # 表头跨页重复
    try:
        table.rows[0].heading = True
    except Exception:
        pass


def export_product_docx(product: dict, note: dict | None, assets: list[dict]) -> bytes:
    doc = Document()
    doc.core_properties.title = product.get("title", "知识付费产品")
    doc.add_heading(product.get("title", "知识付费产品"), 0)
    doc.add_paragraph(
        f"产品类型：{product.get('product_type', '')}    建议售价：¥{product.get('price_suggestion', 0)}"
    )

    # 🔍 [安全/健壮] 2026-08 修复：media_type 缺失时按扩展名兜底，避免漏嵌插图
    def _is_image_asset(a: dict) -> bool:
        if str(a.get("media_type", "")).startswith("image"):
            return True
        name = str(a.get("filename", "")).lower()
        return name.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"))

    image_assets = [a for a in (assets or []) if _is_image_asset(a)]
    inserted: set[int] = set()

    def insert_image(idx: int) -> bool:
        if not (1 <= idx <= len(image_assets)):
            return False
        asset = image_assets[idx - 1]
        path = Path(asset.get("storage_path", ""))
        if not path.is_file():
            return False
        try:
            doc.add_picture(str(path), width=Inches(5.8))
            doc.add_paragraph(f"图 {idx}  {asset.get('filename', '')}", style="Caption")
            inserted.add(idx)
            return True
        except Exception:
            return False

    for block in parse_markdown_blocks(product.get("content", "")):
        kind = block[0]
        if kind == "h":
            level = max(1, min(6, block[1]))
            title = re.sub(r"[*_`~]", "", block[2])
            doc.add_heading(title, level=level)
        elif kind == "bullet":
            for item in block[1]:
                _add_runs(doc.add_paragraph(style="List Bullet"), item)
        elif kind == "ordered":
            for item in block[1]:
                _add_runs(doc.add_paragraph(style="List Number"), item)
        elif kind == "quote":
            _add_runs(doc.add_paragraph(style="Intense Quote"), block[1])
        elif kind == "code":
            lang, code = block[1], block[2]
            if lang:
                doc.add_paragraph(f"语言：{lang}", style="Caption")
            para = doc.add_paragraph(code, style="No Spacing")
            for run in para.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(10)
        elif kind == "hr":
            _add_horizontal_rule(doc)
        elif kind == "table":
            header, rows = block[1], block[2]
            ncols = max(1, len(header))
            # 补齐/截断行，使每行列数与表头一致，避免越界或错位
            norm_rows = []
            for row in rows:
                norm = [(row[c] if c < len(row) else "") for c in range(ncols)]
                norm_rows.append(norm)
            table = doc.add_table(rows=1, cols=ncols)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for c, value in enumerate(header):
                hdr[c].text = ""
                _add_runs(hdr[c].paragraphs[0], value)
                for run in hdr[c].paragraphs[0].runs:
                    run.bold = True
            for row in norm_rows:
                cells = table.add_row().cells
                for c, value in enumerate(row):
                    cells[c].text = ""
                    _add_runs(cells[c].paragraphs[0], value)
            _set_table_widths(table, list(header), norm_rows)
        elif kind == "p":
            text = block[1]
            if IMAGE_MARKER.search(text):
                clean = _strip_markers(text)
                if clean:
                    _add_runs(doc.add_paragraph(), clean)
                for mm in IMAGE_MARKER.finditer(text):
                    if not insert_image(int(mm.group(1))):
                        doc.add_paragraph(_strip_markers(mm.group(0)))
            else:
                _add_runs(doc.add_paragraph(), text)

    usable_assets = [
        (idx, asset)
        for idx, asset in enumerate(image_assets, 1)
        if idx not in inserted and Path(asset.get("storage_path", "")).is_file()
    ]
    if usable_assets:
        doc.add_heading("源笔记插图", level=1)
        doc.add_paragraph("以下插图按其在源 Word 笔记中的先后位置保留，便于交付时继续精修图文布局。")
        for idx, asset in usable_assets:
            if not insert_image(idx):
                doc.add_paragraph(f"[插图无法嵌入：{asset.get('filename', '')}]")
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
