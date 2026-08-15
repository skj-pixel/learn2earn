from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Learn2Earn详细用户手册.md"
OUTPUT = ROOT / "docs" / "Learn2Earn详细用户手册.docx"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x68, 0x72, 0x7D)


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Learn2Earn 用户手册  |  ")
    set_font(run, 9, MUTED)
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc):
    for _ in range(5): doc.add_paragraph()
    kicker = doc.add_paragraph(); kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("PRODUCT GUIDE  |  2026.08"), 10.5, BLUE, True)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("Learn2Earn"), 30, DARK_BLUE, True)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run("图文知识产品工作台\n详细用户手册"), 17, BLUE, True)
    lead = doc.add_paragraph(); lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_before = Pt(20); lead.paragraph_format.space_after = Pt(90)
    set_font(lead.add_run("从 Word 图文笔记、Skills 与后台生成，\n到可编辑 Word 商业交付件的完整操作指南"), 11.5, MUTED)
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(meta.add_run("适用于学习者、知识创作者、课程顾问与个人创业者"), 10, MUTED, italic=True)
    doc.add_page_break()


def add_contents(doc, headings):
    p = doc.add_paragraph("阅读导航", style="Heading 1")
    p.paragraph_format.space_after = Pt(12)
    for i, heading in enumerate(headings, 1):
        item = doc.add_paragraph(style="List Number")
        item.paragraph_format.left_indent = Inches(0.375)
        item.paragraph_format.first_line_indent = Inches(-0.188)
        item.paragraph_format.space_after = Pt(0)
        item.paragraph_format.line_spacing = 1.0
        set_font(item.add_run(heading), 8.5)
    doc.add_page_break()


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part: continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1]); set_font(run, 10, DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); set_font(run, bold=True)
        else:
            set_font(paragraph.add_run(part))


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    configure_styles(doc)
    page_number(section.footer.paragraphs[0])
    add_cover(doc)
    add_contents(doc, headings)

    buffer = []
    def flush():
        if not buffer: return
        p = doc.add_paragraph()
        add_inline(p, "".join(buffer).strip())
        buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped == "---":
            flush(); continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            flush(); doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
        elif stripped.startswith("### "):
            flush(); doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
        elif re.match(r"^[-*] ", stripped):
            flush(); p = doc.add_paragraph(style="List Bullet"); add_inline(p, stripped[2:])
        else:
            buffer.append(stripped + " ")
    flush()
    props = doc.core_properties
    props.title = "Learn2Earn 图文知识产品工作台详细用户手册"
    props.subject = "完整功能、操作流程、术语、测试与故障排查"
    props.author = "Learn2Earn"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
