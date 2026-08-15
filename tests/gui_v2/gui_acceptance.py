import io
import json
import re
import sys
from pathlib import Path

from docx import Document
from PIL import Image
from playwright.sync_api import sync_playwright, expect

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "reports" / "gui_v2"
OUT.mkdir(parents=True, exist_ok=True)
FIXTURE = OUT / "图文课程笔记.docx"
API = "http://127.0.0.1:8123/api"
STEPS = []


def create_fixture():
    image = io.BytesIO()
    Image.new("RGB", (900, 500), "#137c5a").save(image, "PNG")
    image.seek(0)
    doc = Document()
    doc.add_heading("知识产品图文课程", 0)
    doc.add_paragraph("这是一份用于验收 Word 图文导入、表格和插图定位的学习笔记。")
    doc.add_heading("内容规划", 1)
    table = doc.add_table(rows=3, cols=3)
    for cell, text in zip(table.rows[0].cells, ["阶段", "目标", "交付"]): cell.text = text
    for cell, text in zip(table.rows[1].cells, ["专精期", "深入单项能力", "可复用方法"]): cell.text = text
    for cell, text in zip(table.rows[2].cells, ["融合期", "组合多项技能", "完整解决方案"]): cell.text = text
    doc.add_paragraph("下面的插图应当在 Word 导出时保留。")
    doc.add_picture(image)
    doc.add_paragraph("插图之后继续记录产品验收标准和复盘方法。")
    doc.save(FIXTURE)


def mark(page, step_id, description, screenshot=None):
    STEPS.append({"id": step_id, "description": description, "url": page.url})
    if screenshot:
        page.screenshot(path=str(OUT / screenshot), full_page=True)


def main():
    create_fixture()
    errors = []
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport={"width": 1440, "height": 900}, record_video_dir=str(OUT / "video_raw"))
        page = context.new_page()
        page.set_default_timeout(8000)
        page.on("pageerror", lambda exc: errors.append(f"pageerror: {exc}"))
        page.on("console", lambda msg: errors.append(f"console.{msg.type}: {msg.text}") if msg.type == "error" else None)

        auth = page.request.post(f"{API}/auth/signup", data={"email": "codex-editor-v2@example.com", "password": "codex-editor-v2-password"})
        if not auth.ok:
            auth = page.request.post(f"{API}/auth/login", data={"email": "codex-editor-v2@example.com", "password": "codex-editor-v2-password"})
        token = auth.json()["access_token"]
        page.add_init_script(f"localStorage.setItem('learn2earn_access_token', {json.dumps(token)})")
        headers = {"Authorization": f"Bearer {token}"}
        subject_resp = page.request.post(f"{API}/subjects", headers=headers, data={"name": "图文知识产品验收", "icon": "📘", "description": "DOCX、Skills、后台任务验收", "color": "#137c5a"})
        subject = subject_resp.json()

        page.goto(f"http://127.0.0.1:9020/subjects/{subject['id']}/notes/new", wait_until="networkidle")
        expect(page.get_by_text("导入 Word", exact=True)).to_be_visible()
        mark(page, "editor-shell", "打开文档式富文本编辑器", "01_rich_editor.png")

        with page.expect_response(lambda response: "/api/notes/import-docx" in response.url, timeout=30000) as import_info:
            page.locator('input[type="file"][accept*="docx"]').set_input_files(str(FIXTURE))
        import_response = import_info.value
        if not import_response.ok:
            raise AssertionError(f"DOCX import failed ({import_response.status}): {import_response.text()}")
        page.wait_for_url(re.compile(r"/notes/\d+$"), timeout=30000)
        page.wait_for_timeout(2500)
        page.screenshot(path=str(OUT / "debug_after_import.png"), full_page=True)
        print("after import", page.url, "tables", page.locator('.knowledge-editor table').count(), "images", page.locator('.knowledge-editor img').count())
        expect(page.locator('.knowledge-editor table')).to_be_visible(timeout=10000)
        expect(page.locator('.knowledge-editor img')).to_be_visible(timeout=10000)
        mark(page, "docx-import", "导入 Word 并识别标题、表格和插图", "02_docx_import.png")

        page.get_by_title("插入 3×3 表格").click()
        page.get_by_title("撤销").click()
        page.get_by_title("重做").click()
        page.get_by_role("button", name="保存").click()
        expect(page.get_by_text("笔记已保存")).to_be_visible()
        note_id = int(page.url.rstrip('/').split('/')[-1])
        mark(page, "editor-tools", "操作表格、撤销重做并保存富文本", "03_editor_tools.png")

        product_resp = page.request.post(f"{API}/products", headers=headers, data={
            "title": "图文课程商业化验收产品", "product_type": "software_tutorial",
            "content": "# 软件使用教程\n\n| 功能 | 用户收益 |\n|---|---|\n| Word导入 | 保留图文结构 |\n| 后台任务 | 切换页面不中断 |\n\n[插图 1: word-image-1.png]\n\n## 分销说明\n明确功能、收益和佣金披露。",
            "subject_id": subject["id"], "note_id": note_id, "price_suggestion": 99,
            "platform_suggestion": ["公众号", "小红书"], "keywords": ["Word", "教程"], "status": "draft",
        })
        product = product_resp.json()

        page.goto(f"http://127.0.0.1:9020/products/{product['id']}", wait_until="networkidle")
        expect(page.locator('.markdown-body table')).to_be_visible()
        mark(page, "markdown-table", "正确渲染 GFM Markdown 表格", "04_markdown_table.png")
        with page.expect_download(timeout=15000) as download_info:
            page.get_by_role("button", name="📄 导出 Word").click()
        exported = OUT / "exported_product.docx"
        download_info.value.save_as(exported)
        exported_doc = Document(exported)
        assert exported_doc.inline_shapes, "exported Word should contain the source image"
        mark(page, "word-export", "仅保留 Word 导出并验证内嵌源笔记插图")

        page.goto("http://127.0.0.1:9020/skills", wait_until="networkidle")
        page.get_by_role("button", name="导入知识付费包").click()
        expect(page.get_by_text("从知识付费包安装", exact=False)).to_be_visible(timeout=60000)
        expect(page.get_by_role("heading", name="Skills 仓库", exact=True)).to_be_visible()
        mark(page, "skills-store", "一键安装本机知识付费 Skills 包", "05_skills_store.png")

        page.goto(f"http://127.0.0.1:9020/notes/{note_id}/generate", wait_until="networkidle")
        expect(page.get_by_text("生成策略", exact=True)).to_be_visible()
        skill_option = page.locator('label:has(input[type="checkbox"])').first
        if skill_option.count(): skill_option.click()
        page.get_by_text("分块生成", exact=True).click()
        page.get_by_text("示例驱动", exact=True).click()
        mark(page, "generation-options", "选择 Skill、分块算法和质量技术", "06_generation_options.png")
        page.get_by_role("button", name="生成", exact=True).first.click()
        page.wait_for_url("**/tasks", timeout=15000)
        expect(page.get_by_text("生成任务", exact=True).first).to_be_visible()
        mark(page, "background-task", "提交后台任务后离开生成页并持续查看进度", "07_background_tasks.png")

        page.goto(f"http://127.0.0.1:9020/subjects/{subject['id']}/notes", wait_until="networkidle")
        page.get_by_placeholder("按名称搜索笔记").fill("图文课程")
        page.locator('.resource-sort select').select_option('size')
        expect(page.get_by_text("图文课程笔记", exact=False)).to_be_visible()
        mark(page, "note-search-sort", "按名称搜索笔记并按大小排序", "08_note_search_sort.png")

        page.goto("http://127.0.0.1:9020/products", wait_until="networkidle")
        page.get_by_placeholder("按名称搜索产品").fill("商业化验收")
        page.locator('.resource-sort select').select_option('name')
        expect(page.get_by_text("图文课程商业化验收产品", exact=True).first).to_be_visible()
        mark(page, "product-search-sort", "按名称搜索产品并排序", "09_product_search_sort.png")

        page.goto("http://127.0.0.1:9020/subjects", wait_until="networkidle")
        page.get_by_placeholder("按名称搜索科目").fill("图文知识")
        page.locator('.resource-sort select').select_option('updated')
        expect(page.get_by_text("图文知识产品验收", exact=True).first).to_be_visible()
        mark(page, "subject-search-sort", "按名称搜索科目并按最后修改时间排序", "10_subject_search_sort.png")

        context.close()

    ignored = ("favicon.ico",)
    actionable = [item for item in errors if not any(token in item for token in ignored)]
    (OUT / "steps.json").write_text(json.dumps(STEPS, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT / "browser_errors.json").write_text(json.dumps(actionable, ensure_ascii=False, indent=2), encoding="utf-8")
    coverage = "# GUI 功能覆盖\n\n" + "\n".join(f"- [x] `{step['id']}`：{step['description']}" for step in STEPS)
    coverage += f"\n\n浏览器错误：{len(actionable)}\n"
    (OUT / "coverage_report.md").write_text(coverage, encoding="utf-8")
    if actionable:
        raise AssertionError("browser errors: " + "; ".join(actionable))
    print(f"GUI PASS: {len(STEPS)} steps; artifacts={OUT}")


if __name__ == "__main__":
    if hasattr(sys.stdout, "reconfigure"): sys.stdout.reconfigure(encoding="utf-8")
    main()
